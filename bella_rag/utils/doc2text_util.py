import io
import subprocess
import tempfile
from io import IOBase
from typing import Optional

from init.settings import user_logger

logger = user_logger


def convert_docx_to_text(contents: bytes) -> Optional[str]:
    """
    将 DOCX 文件直接转换为文本
    使用轻量级的 python-docx 库
    
    Args:
        contents: DOCX 文件内容
        
    Returns:
        str: 提取的文本内容，失败时返回 None
    """
    try:
        from docx import Document

        # 创建输入流
        docx_stream = io.BytesIO(contents)
        docx_stream.seek(0)

        # 加载 DOCX 文档
        doc = Document(docx_stream)

        # 提取所有段落文本
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 跳过空段落
                full_text.append(paragraph.text)

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))

        text_content = "\n".join(full_text)
        logger.info(f'DOCX text extraction successful: {len(text_content)} characters')
        return text_content

    except ImportError:
        logger.error("python-docx not available. Please install: pip install python-docx")
        return None
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {str(e)}")
        return None


def convert_doc_to_text(contents: bytes) -> Optional[str]:
    """
    将 DOC 文件转换为文本
    使用 antiword (轻量级 DOC 文本提取工具)
    
    Args:
        contents: DOC 文件内容
        
    Returns:
        str: 提取的文本内容，失败时返回 None
    """
    return _convert_doc_with_antiword(contents)


def _convert_doc_with_antiword(contents: bytes) -> Optional[str]:
    """使用 antiword 提取 DOC 文本"""
    try:
        with tempfile.NamedTemporaryFile(suffix='.doc') as temp_file:
            temp_file.write(contents)
            temp_file.flush()

            result = subprocess.run(
                ['antiword', temp_file.name],
                capture_output=True,
                timeout=30,
                text=True
            )

            if result.returncode == 0:
                text_content = result.stdout.strip()
                logger.info(f'DOC text extraction successful: {len(text_content)} characters')
                return text_content
            else:
                logger.error(f"antiword failed: {result.stderr}")
                return None

    except FileNotFoundError:
        logger.error("antiword not found. Please install: brew install antiword")
        return None
    except Exception as e:
        logger.error(f"DOC text extraction failed: {str(e)}")
        return None


def convert_docx_to_text_stream(docx_stream: IOBase) -> Optional[str]:
    """
    从 DOCX 流中提取文本
    
    Args:
        docx_stream: DOCX 文件流
        
    Returns:
        str: 提取的文本内容
    """
    docx_stream.seek(0)
    contents = docx_stream.read()
    return convert_docx_to_text(contents)


def convert_doc_to_text_stream(doc_stream: IOBase) -> Optional[str]:
    """
    从 DOC 流中提取文本
    
    Args:
        doc_stream: DOC 文件流
        
    Returns:
        str: 提取的文本内容
    """
    doc_stream.seek(0)
    contents = doc_stream.read()
    return convert_doc_to_text(contents)
